from flask import Flask, jsonify, request
from flask_cors import CORS
from config import Config
from models import db, User, Task, Trip, ItineraryItem
from datetime import datetime, timedelta
import os
import re
import base64
from PIL import Image
import pytesseract
from io import BytesIO
import json

import requests

DEEPSEEK_AVAILABLE = True

TESSERACT_PATHS = [
    r'D:\TesseractOCR\tesseract.exe',
    r'D:\TesseractOCR\bin\tesseract.exe',
    r'D:\Tesseract-OCR\tesseract.exe',
    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
    r'D:\Tesseract\tesseract.exe'
]

print("[DEBUG] Searching for Tesseract executable...")
tesseract_found = False
tessdata_dir = None

for tess_path in TESSERACT_PATHS:
    if os.path.exists(tess_path):
        pytesseract.pytesseract.tesseract_cmd = tess_path
        print(f"[INFO] Tesseract found at: {tess_path}")
        
        tess_dir = os.path.dirname(tess_path)
        potential_tessdata = os.path.join(tess_dir, 'tessdata')
        if os.path.exists(potential_tessdata):
            tessdata_dir = potential_tessdata
            print(f"[INFO] Found tessdata directory: {tessdata_dir}")
            os.environ['TESSDATA_PREFIX'] = tessdata_dir
        
        tesseract_found = True
        break
else:
    print("[WARNING] Tesseract not found in common paths.")
    print("[WARNING] Checked paths:")
    for p in TESSERACT_PATHS:
        print(f"         {p} - {'EXISTS' if os.path.exists(p) else 'NOT FOUND'}")
        
    import subprocess
    try:
        result = subprocess.run(['where', 'tesseract'], capture_output=True, text=True)
        if result.returncode == 0 and result.stdout.strip():
            found_path = result.stdout.strip().split('\n')[0]
            pytesseract.pytesseract.tesseract_cmd = found_path
            print(f"[INFO] Tesseract found via 'where' command: {found_path}")
            
            tess_dir = os.path.dirname(found_path)
            potential_tessdata = os.path.join(tess_dir, 'tessdata')
            if os.path.exists(potential_tessdata):
                tessdata_dir = potential_tessdata
                print(f"[INFO] Found tessdata directory: {tessdata_dir}")
                os.environ['TESSDATA_PREFIX'] = tessdata_dir
            
            tesseract_found = True
        else:
            print("[WARNING] Tesseract not found via 'where' command")
    except Exception as e:
        print(f"[ERROR] Failed to run 'where' command: {e}")

if tesseract_found:
    print("[INFO] Tesseract OCR is ready")
    
    if tessdata_dir:
        chi_sim_path = os.path.join(tessdata_dir, 'chi_sim.traineddata')
        eng_path = os.path.join(tessdata_dir, 'eng.traineddata')
        print(f"[DEBUG] Checking language files:")
        print(f"        chi_sim.traineddata: {'EXISTS' if os.path.exists(chi_sim_path) else 'MISSING'}")
        print(f"        eng.traineddata: {'EXISTS' if os.path.exists(eng_path) else 'MISSING'}")
else:
    print("[WARNING] Tesseract OCR is not available")

app = Flask(__name__)
app.config.from_object(Config)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
CORS(app)
db.init_app(app)

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok', 'message': 'AI Task Management System API is running'})

@app.route('/api/users', methods=['POST'])
def create_user():
    try:
        data = request.get_json()
        if not data or 'username' not in data or 'email' not in data or 'password' not in data:
            return jsonify({'error': 'Missing required fields'}), 400
        
        if User.query.filter_by(username=data['username']).first():
            return jsonify({'error': 'Username already exists'}), 400
        
        if User.query.filter_by(email=data['email']).first():
            return jsonify({'error': 'Email already exists'}), 400
        
        new_user = User(
            username=data['username'],
            email=data['email'],
            password_hash=data['password']
        )
        db.session.add(new_user)
        db.session.commit()
        
        return jsonify({'message': 'User created successfully', 'user_id': new_user.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/users/login', methods=['POST'])
def login_user():
    try:
        data = request.get_json()
        if not data or 'username' not in data or 'password' not in data:
            return jsonify({'error': 'Missing username or password'}), 400
        
        user = User.query.filter_by(username=data['username']).first()
        if user and user.password_hash == data['password']:
            return jsonify({
                'message': 'Login successful',
                'user': {
                    'id': user.id,
                    'username': user.username,
                    'email': user.email
                }
            }), 200
        else:
            return jsonify({'error': 'Invalid username or password'}), 401
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# 任务API路由
@app.route('/api/tasks', methods=['GET'])
def get_tasks():
    try:
        user_id = request.args.get('user_id')
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 10))
        
        if not user_id:
            return jsonify({'error': 'User ID is required'}), 400
        
        pagination = Task.query.filter_by(user_id=user_id).order_by(Task.start_time).paginate(page=page, per_page=per_page, error_out=False)
        result = []
        for t in pagination.items:
            result.append({
                'id': t.id,
                'title': t.title,
                'description': t.description,
                'start_time': t.start_time.isoformat(),
                'end_time': t.end_time.isoformat(),
                'location': t.location,
                'priority': t.priority,
                'status': t.status,
                'category': t.category,
                'created_at': t.created_at.isoformat()
            })
        return jsonify({
            'items': result,
            'total': pagination.total,
            'page': pagination.page,
            'per_page': pagination.per_page,
            'pages': pagination.pages
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/tasks/batch', methods=['DELETE'])
def batch_delete_tasks():
    try:
        data = request.get_json()
        if not data or 'ids' not in data:
            return jsonify({'error': 'Missing required fields'}), 400
        
        ids = data['ids']
        if not isinstance(ids, list) or len(ids) == 0:
            return jsonify({'error': 'Invalid IDs list'}), 400
        
        Task.query.filter(Task.id.in_(ids)).delete(synchronize_session=False)
        db.session.commit()
        
        return jsonify({'message': f'Deleted {len(ids)} tasks'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/tasks', methods=['POST'])
def create_task():
    try:
        data = request.get_json()
        required_fields = ['user_id', 'title', 'start_time', 'end_time']
        if not data or any(f not in data for f in required_fields):
            return jsonify({'error': 'Missing required fields'}), 400
        
        new_task = Task(
            user_id=data['user_id'],
            title=data['title'],
            description=data.get('description'),
            start_time=datetime.fromisoformat(data['start_time']),
            end_time=datetime.fromisoformat(data['end_time']),
            location=data.get('location'),
            priority=data.get('priority', 2),
            status=data.get('status', 'pending'),
            category=data.get('category', 'personal')
        )
        db.session.add(new_task)
        db.session.commit()
        
        return jsonify({'message': 'Task created successfully', 'task_id': new_task.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/tasks/<int:id>', methods=['PUT'])
def update_task(id):
    try:
        task = Task.query.get(id)
        if not task:
            return jsonify({'error': 'Task not found'}), 404
        
        data = request.get_json()
        if 'title' in data:
            task.title = data['title']
        if 'description' in data:
            task.description = data['description']
        if 'start_time' in data:
            task.start_time = datetime.fromisoformat(data['start_time'])
        if 'end_time' in data:
            task.end_time = datetime.fromisoformat(data['end_time'])
        if 'location' in data:
            task.location = data['location']
        if 'priority' in data:
            task.priority = data['priority']
        if 'status' in data:
            task.status = data['status']
        if 'category' in data:
            task.category = data['category']
        
        db.session.commit()
        return jsonify({'message': 'Task updated successfully'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/tasks/<int:id>', methods=['DELETE'])
def delete_task(id):
    try:
        task = Task.query.get(id)
        if not task:
            print(f"[DEBUG] Task with id {id} not found")
            return jsonify({'error': 'Task not found'}), 404
        
        print(f"[DEBUG] Deleting task: {task.title} (id: {id})")
        db.session.delete(task)
        db.session.commit()
        print(f"[DEBUG] Task deleted successfully")
        
        remaining_tasks = Task.query.filter_by(user_id=task.user_id).count()
        print(f"[DEBUG] Remaining tasks: {remaining_tasks}")
        
        return jsonify({'message': 'Task deleted successfully'}), 200
    except Exception as e:
        db.session.rollback()
        print(f"[DEBUG] Error deleting task: {str(e)}")
        return jsonify({'error': str(e)}), 500

# 行程API路由
@app.route('/api/trips', methods=['GET'])
def get_trips():
    try:
        user_id = request.args.get('user_id')
        if not user_id:
            return jsonify({'error': 'User ID is required'}), 400
        
        trips = Trip.query.filter_by(user_id=user_id).order_by(Trip.start_date).all()
        result = []
        for t in trips:
            result.append({
                'id': t.id,
                'title': t.title,
                'description': t.description,
                'destination': t.destination,
                'start_date': t.start_date.isoformat(),
                'end_date': t.end_date.isoformat(),
                'budget': t.budget,
                'status': t.status,
                'created_at': t.created_at.isoformat()
            })
        return jsonify(result), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/trips', methods=['POST'])
def create_trip():
    try:
        data = request.get_json()
        required_fields = ['user_id', 'title', 'destination', 'start_date', 'end_date']
        if not data or any(f not in data for f in required_fields):
            return jsonify({'error': 'Missing required fields'}), 400
        
        new_trip = Trip(
            user_id=data['user_id'],
            title=data['title'],
            description=data.get('description'),
            destination=data['destination'],
            start_date=datetime.fromisoformat(data['start_date']).date(),
            end_date=datetime.fromisoformat(data['end_date']).date(),
            budget=data.get('budget'),
            status=data.get('status', 'planned')
        )
        db.session.add(new_trip)
        db.session.commit()
        
        return jsonify({'message': 'Trip created successfully', 'trip_id': new_trip.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/trips/<int:id>', methods=['PUT'])
def update_trip(id):
    try:
        trip = Trip.query.get(id)
        if not trip:
            return jsonify({'error': 'Trip not found'}), 404
        
        data = request.get_json()
        if 'title' in data:
            trip.title = data['title']
        if 'description' in data:
            trip.description = data['description']
        if 'destination' in data:
            trip.destination = data['destination']
        if 'start_date' in data:
            trip.start_date = datetime.fromisoformat(data['start_date']).date()
        if 'end_date' in data:
            trip.end_date = datetime.fromisoformat(data['end_date']).date()
        if 'budget' in data:
            trip.budget = data['budget']
        if 'status' in data:
            trip.status = data['status']
        
        db.session.commit()
        return jsonify({'message': 'Trip updated successfully'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/trips/batch', methods=['DELETE'])
def batch_delete_trips():
    try:
        data = request.get_json()
        if not data or 'ids' not in data:
            return jsonify({'error': 'Missing required fields'}), 400

        ids = data['ids']
        if not isinstance(ids, list) or len(ids) == 0:
            return jsonify({'error': 'Invalid IDs list'}), 400

        # 先删除所有关联的行程项目
        ItineraryItem.query.filter(ItineraryItem.trip_id.in_(ids)).delete(synchronize_session=False)

        # 再删除行程
        Trip.query.filter(Trip.id.in_(ids)).delete(synchronize_session=False)
        db.session.commit()

        return jsonify({'message': f'Deleted {len(ids)} trips'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/trips/<int:id>', methods=['DELETE'])
def delete_trip(id):
    try:
        trip = Trip.query.get(id)
        if not trip:
            return jsonify({'error': 'Trip not found'}), 404

        # 先删除关联的行程项目
        ItineraryItem.query.filter_by(trip_id=id).delete()

        db.session.delete(trip)
        db.session.commit()
        print(f"[DEBUG] Trip {id} and its itinerary items deleted successfully")
        return jsonify({'message': 'Trip deleted successfully'}), 200
    except Exception as e:
        db.session.rollback()
        print(f"[ERROR] Delete trip failed: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/trips/<int:id>/itinerary', methods=['GET'])
def get_itinerary(id):
    try:
        items = ItineraryItem.query.filter_by(trip_id=id).order_by(ItineraryItem.day_number).all()
        result = []
        for item in items:
            result.append({
                'id': item.id,
                'day_number': item.day_number,
                'title': item.title,
                'description': item.description,
                'location': item.location,
                'time_slot': item.time_slot
            })
        return jsonify(result), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/trips/<int:id>/itinerary', methods=['POST'])
def add_itinerary_item(id):
    try:
        data = request.get_json()
        required_fields = ['day_number', 'title']
        if not data or any(f not in data for f in required_fields):
            return jsonify({'error': 'Missing required fields'}), 400
        
        new_item = ItineraryItem(
            trip_id=id,
            day_number=data['day_number'],
            title=data['title'],
            description=data.get('description'),
            location=data.get('location'),
            time_slot=data.get('time_slot')
        )
        db.session.add(new_item)
        db.session.commit()
        
        return jsonify({'message': 'Itinerary item added successfully', 'item_id': new_item.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# AI规划相关函数
def get_user_tasks(user_id):
    try:
        tasks = Task.query.filter_by(user_id=user_id).order_by(Task.start_time).all()
        return [{
            'id': t.id,
            'title': t.title,
            'description': t.description,
            'start_time': t.start_time.isoformat(),
            'end_time': t.end_time.isoformat(),
            'location': t.location,
            'priority': t.priority,
            'status': t.status,
            'category': t.category
        } for t in tasks]
    except:
        return []

def get_user_trips(user_id):
    try:
        trips = Trip.query.filter_by(user_id=user_id).order_by(Trip.start_date).all()
        return [{
            'id': t.id,
            'title': t.title,
            'description': t.description,
            'destination': t.destination,
            'start_date': t.start_date.isoformat(),
            'end_date': t.end_date.isoformat(),
            'budget': t.budget,
            'status': t.status
        } for t in trips]
    except:
        return []

@app.route('/api/ai/plan', methods=['POST'])
def ai_plan():
    try:
        data = request.get_json()
        if not data or 'user_id' not in data or 'prompt' not in data:
            return jsonify({'error': '缺少必要参数'}), 400
        
        user_id = data['user_id']
        prompt = data['prompt']
        
        existing_tasks = get_user_tasks(user_id)
        existing_trips = get_user_trips(user_id)
        
        analysis = analyze_schedule_for_plan(prompt, existing_tasks, existing_trips)
        
        suggestions = []
        ai_source = 'fallback'
        
        if DEEPSEEK_AVAILABLE:
            try:
                suggestions = call_deepseek_api(prompt, existing_tasks, existing_trips)
                ai_source = 'deepseek'
            except Exception as deepseek_error:
                suggestions = generate_ai_suggestions(prompt, existing_tasks, existing_trips)
        else:
            suggestions = generate_ai_suggestions(prompt, existing_tasks, existing_trips)
        
        # 日期修正：确保所有旅行计划的日期符合用户指定的月份且是未来日期
        today = datetime.now()
        month_match = re.search(r'(\d+)月(份)?', prompt)
        if not month_match:
            month_match = re.search(r'([一二三四五六七八九十]+)月(份)?', prompt)
        
        if month_match:
            month_name = month_match.group(1)
            if month_name.isdigit():
                target_month = int(month_name)
            else:
                target_month = {'一':1, '二':2, '三':3, '四':4, '五':5, '六':6, 
                            '七':7, '八':8, '九':9, '十':10, '十一':11, '十二':12}.get(month_name)
            
            if target_month and 1 <= target_month <= 12 and suggestions:
                for s in suggestions:
                    if s.get('type') == 'trip' and 'start_date' in s:
                        # 获取AI返回的开始日期
                        ai_start_date = s.get('start_date', '')
                        try:
                            start_dt = datetime.strptime(ai_start_date, '%Y-%m-%d')
                        except:
                            start_dt = None
                        
                        # 确定年份
                        target_year = today.year
                        if target_month < today.month:
                            target_year += 1
                        
                        # 如果AI返回的日期在目标月份内且是未来日期，使用AI的日期
                        if start_dt and start_dt.month == target_month and start_dt.year == target_year and start_dt.date() >= today.date():
                            pass  # 保持AI返回的日期
                        else:
                            # 否则计算未来的日期
                            # 如果目标月份是当前月份，从明天开始找合适的日期
                            if target_month == today.month and target_year == today.year:
                                start_day = today.day + 1
                                if start_day > 28:
                                    start_day = today.day + 3
                            else:
                                start_day = 1
                            
                            # 确保日期有效
                            try:
                                start_date = datetime(target_year, target_month, min(start_day, 28))
                                s['start_date'] = start_date.strftime('%Y-%m-%d')
                                
                                # 计算结束日期（假设4天行程）
                                end_date = start_date + timedelta(days=3)
                                s['end_date'] = end_date.strftime('%Y-%m-%d')
                            except:
                                # 如果日期无效，使用默认值
                                s['start_date'] = f"{target_year}-{target_month:02d}-01"
                                s['end_date'] = f"{target_year}-{target_month:02d}-04"
        
        if not suggestions or len(suggestions) == 0:
            suggestions = generate_default_suggestions(prompt)
        else:
            filtered_suggestions = []
            conflict_keywords = ['冲突', '调整时间', '重新安排', '修改时间', '时间冲突']
            for s in suggestions:
                title = s.get('title', '')
                description = s.get('description', '')
                # 过滤掉冲突解决相关的任务
                is_conflict_task = any(keyword in title or keyword in description for keyword in conflict_keywords)
                if not is_conflict_task:
                    filtered_suggestions.append(s)
            
            if not filtered_suggestions:
                filtered_suggestions = generate_default_suggestions(prompt)
            suggestions = filtered_suggestions
        
        return jsonify({
            'message': 'AI规划完成',
            'suggestions': suggestions,
            'analysis': analysis,
            'ai_source': ai_source
        }), 200
    except Exception as e:
        print(f"[AI PLAN] Error: {str(e)}")
        return jsonify({
            'error': 'AI规划失败，请稍后重试',
            'suggestions': generate_default_suggestions(prompt),
            'analysis': {'summary': '分析过程中出现错误', 'has_conflicts': False, 'conflicts': [], 'suggestions_count': 0}
        }), 200

def generate_default_suggestions(prompt):
    suggestions = []
    today = datetime.now()
    
    suggestions.append({
        'type': 'task',
        'title': '日常安排',
        'description': '根据您的需求，建议合理安排日常时间',
        'start_time': today.replace(hour=9, minute=0).isoformat(),
        'end_time': today.replace(hour=10, minute=0).isoformat(),
        'location': '办公室',
        'priority': 2,
        'category': 'work'
    })
    
    return suggestions

def analyze_schedule_for_plan(prompt, existing_tasks=None, existing_trips=None):
    today = datetime.now()
    existing_tasks = existing_tasks or []
    existing_trips = existing_trips or []
    
    analysis = {
        'summary': '',
        'has_conflicts': False,
        'conflicts': [],
        'busy_periods': [],
        'free_periods': [],
        'upcoming_events': [],
        'suggestions_count': len(existing_tasks) + len(existing_trips),
        'priority_summary': {
            'high': 0,
            'medium': 0,
            'low': 0
        },
        'category_summary': {},
        'recommendations': []
    }
    
    all_events = []
    
    for t in existing_tasks:
        try:
            start = datetime.fromisoformat(t['start_time'].replace('Z', '+00:00'))
            end = datetime.fromisoformat(t['end_time'].replace('Z', '+00:00'))
            all_events.append({
                'start': start,
                'end': end,
                'title': t['title'],
                'type': 'task',
                'priority': t.get('priority', 2),
                'category': t.get('category', 'personal')
            })
            
            priority = t.get('priority', 2)
            if priority == 1:
                analysis['priority_summary']['high'] += 1
            elif priority == 2:
                analysis['priority_summary']['medium'] += 1
            else:
                analysis['priority_summary']['low'] += 1
            
            category = t.get('category', 'personal')
            analysis['category_summary'][category] = analysis['category_summary'].get(category, 0) + 1
        except:
            pass
    
    for t in existing_trips:
        try:
            start = datetime.strptime(t['start_date'], '%Y-%m-%d')
            end = datetime.strptime(t['end_date'], '%Y-%m-%d')
            all_events.append({
                'start': start,
                'end': end,
                'title': t['title'],
                'type': 'trip',
                'priority': 1,
                'category': 'travel'
            })
            analysis['priority_summary']['high'] += 1
            analysis['category_summary']['travel'] = analysis['category_summary'].get('travel', 0) + 1
        except:
            pass
    
    all_events.sort(key=lambda x: x['start'])
    
    for i in range(len(all_events) - 1):
        if all_events[i]['end'] > all_events[i+1]['start']:
            analysis['has_conflicts'] = True
            
            event1_priority = all_events[i].get('priority', 2)
            event2_priority = all_events[i+1].get('priority', 2)
            
            priority_map = {1: '高优先级', 2: '中优先级', 3: '低优先级'}
            type_map = {'task': '任务', 'trip': '行程'}
            
            if event1_priority < event2_priority:
                recommended_event = 'event1'
                recommended_title = all_events[i]['title']
                recommended_reason = f'"{all_events[i]["title"]}"为{priority_map[event1_priority]}，高于"{all_events[i+1]["title"]}"的{priority_map[event2_priority]}，建议优先执行'
            elif event2_priority < event1_priority:
                recommended_event = 'event2'
                recommended_title = all_events[i+1]['title']
                recommended_reason = f'"{all_events[i+1]["title"]}"为{priority_map[event2_priority]}，高于"{all_events[i]["title"]}"的{priority_map[event1_priority]}，建议优先执行'
            else:
                if all_events[i]['type'] == 'trip' and all_events[i+1]['type'] == 'task':
                    recommended_event = 'event2'
                    recommended_title = all_events[i+1]['title']
                    recommended_reason = f'两者优先级相同，但"{all_events[i+1]["title"]}"是{type_map[all_events[i+1]["type"]]}，通常任务的时间更刚性，建议优先执行'
                elif all_events[i]['type'] == 'task' and all_events[i+1]['type'] == 'trip':
                    recommended_event = 'event1'
                    recommended_title = all_events[i]['title']
                    recommended_reason = f'两者优先级相同，但"{all_events[i]["title"]}"是{type_map[all_events[i]["type"]]}，通常任务的时间更刚性，建议优先执行'
                else:
                    recommended_event = 'event1'
                    recommended_title = all_events[i]['title']
                    recommended_reason = f'两者优先级相同，"{all_events[i]["title"]}"时间更早，建议优先执行'
            
            analysis['conflicts'].append({
                'event1': {
                    'title': all_events[i]['title'],
                    'start': all_events[i]['start'].isoformat(),
                    'end': all_events[i]['end'].isoformat(),
                    'type': all_events[i]['type'],
                    'priority': event1_priority,
                    'priority_label': priority_map[event1_priority]
                },
                'event2': {
                    'title': all_events[i+1]['title'],
                    'start': all_events[i+1]['start'].isoformat(),
                    'end': all_events[i+1]['end'].isoformat(),
                    'type': all_events[i+1]['type'],
                    'priority': event2_priority,
                    'priority_label': priority_map[event2_priority]
                },
                'conflict_type': 'time_overlap',
                'recommended_event': recommended_event,
                'recommended_title': recommended_title,
                'recommendation': recommended_reason
            })
    
    analysis['upcoming_events'] = [
        {
            'title': e['title'],
            'start': e['start'].isoformat(),
            'type': e['type']
        } for e in all_events if e['start'] >= today and e['start'] <= today + timedelta(days=7)
    ]
    
    work_hours_start = today.replace(hour=9, minute=0, second=0)
    work_hours_end = today.replace(hour=18, minute=0, second=0)
    
    for day_offset in range(7):
        day_start = (today + timedelta(days=day_offset)).replace(hour=0, minute=0, second=0)
        day_end = (today + timedelta(days=day_offset + 1)).replace(hour=0, minute=0, second=0)
        
        day_events = [e for e in all_events if e['start'] < day_end and e['end'] > day_start]
        day_events.sort(key=lambda x: x['start'])
        
        occupied_minutes = 0
        for e in day_events:
            overlap_start = max(e['start'], day_start)
            overlap_end = min(e['end'], day_end)
            occupied_minutes += int((overlap_end - overlap_start).total_seconds() / 60)
        
        analysis['busy_periods'].append({
            'date': day_start.date().isoformat(),
            'busy_minutes': occupied_minutes,
            'total_minutes': 1440,
            'busy_percentage': round(occupied_minutes / 1440 * 100, 1),
            'event_count': len(day_events)
        })
    
    if analysis['has_conflicts']:
        analysis['recommendations'].append('检测到时间冲突，请及时调整相关任务的时间安排')
    
    total_busy = sum(p['busy_minutes'] for p in analysis['busy_periods'][:5])
    avg_busy_percent = total_busy / (5 * 1440) * 100
    
    if avg_busy_percent > 70:
        analysis['recommendations'].append('近期日程安排较满，建议合理分配时间，避免过度劳累')
    elif avg_busy_percent < 30:
        analysis['recommendations'].append('近期日程较为空闲，可以安排更多任务或活动')
    
    high_priority_count = analysis['priority_summary']['high']
    if high_priority_count > 5:
        analysis['recommendations'].append('高优先级任务较多，建议优先处理重要事项')
    
    if 'travel' in analysis['category_summary'] and analysis['category_summary']['travel'] > 0:
        analysis['recommendations'].append('近期有出行计划，请提前做好准备')
    
    encouragement_messages = [
        '您做得很棒！合理规划时间是成功的第一步，继续保持！',
        '劳逸结合才能走得更远，记得给自己留出休息的时间哦~',
        '每一个小目标的完成都是向梦想迈进的一步，加油！',
        '时间管理是一门艺术，您正在不断精进，相信自己！',
        '忙忙碌碌的日子里，别忘了停下脚步欣赏身边的美好~',
        '您的努力终将得到回报，坚持就是胜利！',
        '合理安排时间，让生活更有条理，让自己更加从容~'
    ]
    
    import random
    analysis['encouragement'] = random.choice(encouragement_messages)
    
    summary_parts = []
    total_events = analysis['suggestions_count']
    
    if total_events == 0:
        summary_parts.append('您目前没有任何任务安排')
    else:
        summary_parts.append(f'您当前共有 {total_events} 个任务/行程安排')
        
        if analysis['priority_summary']['high'] > 0:
            summary_parts.append(f'其中 {analysis["priority_summary"]["high"]} 个高优先级任务')
        
        if analysis['has_conflicts']:
            summary_parts.append(f'检测到 {len(analysis["conflicts"])} 个时间冲突')
        
        if analysis['upcoming_events']:
            summary_parts.append(f'未来7天有 {len(analysis["upcoming_events"])} 个待办事项')
    
    total_busy = sum(p['busy_minutes'] for p in analysis['busy_periods'][:7])
    avg_busy_percent = total_busy / (7 * 1440) * 100
    
    if avg_busy_percent > 75:
        summary_parts.append('近期日程非常紧凑，请注意合理安排休息时间，保持高效工作的同时也要照顾好自己')
    elif avg_busy_percent > 50:
        summary_parts.append('近期日程安排适中，继续保持良好的时间管理习惯')
    else:
        summary_parts.append('近期日程较为轻松，可以考虑安排一些新的任务或活动')
    
    if analysis['has_conflicts']:
        summary_parts.append('建议您检查并调整冲突任务的时间，合理规划可以让生活更加从容')
    
    summary_parts.append(f'🌟 {analysis["encouragement"]}')
    
    if re.search(r'(旅行|旅游|出差|vacation|trip)', prompt, re.IGNORECASE):
        summary_parts.append('根据您的需求，系统已分析您的时间安排，帮助您规划出行计划')
    
    if re.search(r'(会议|meeting|讨论|沟通)', prompt, re.IGNORECASE):
        summary_parts.append('根据您的需求，系统已分析您的时间安排，帮助您安排会议')
    
    analysis['summary'] = '；'.join(summary_parts)
    
    return analysis

def call_deepseek_api(prompt, existing_tasks=None, existing_trips=None):
    try:
        today = datetime.now()
        
        existing_data = ""
        if existing_tasks and len(existing_tasks) > 0:
            existing_data += f"\n用户现有任务（{len(existing_tasks)}个）：\n"
            for t in existing_tasks[:5]:
                existing_data += f"- {t.get('title', '')}: {t.get('start_time', '')} - {t.get('location', '')}\n"
        
        if existing_trips and len(existing_trips) > 0:
            existing_data += f"\n用户现有行程（{len(existing_trips)}个）：\n"
            for t in existing_trips[:3]:
                existing_data += f"- {t.get('title', '')}: {t.get('destination', '')} ({t.get('start_date', '')} - {t.get('end_date', '')})\n"
        
        system_prompt = """
        你是一个专业的智能任务规划助手，擅长根据用户的现有任务安排和需求，生成合理、可行的任务/行程建议。

        核心能力：
        1. 智能分析用户需求，理解用户想要做什么
        2. 参考用户现有的任务安排，避免时间冲突
        3. 生成结构化、可执行的建议
        4. 根据需求类型智能分类（任务/行程）

        重要要求：
        1. 不要生成"冲突解决"、"时间冲突修改"等相关的任务建议
        2. 不要生成"调整时间"、"重新安排"等类型的任务
        3. 如果检测到时间冲突，只需在分析报告中说明，不需要生成冲突解决任务
        4. 只生成实际的行程规划或有用的任务建议

        输出格式要求：
        请输出一个JSON数组，每个元素包含以下字段：
        - type: 类型，可选值: "task"（任务）, "trip"（行程）
        - title: 标题（简洁明了，不超过20字）
        - description: 描述（详细说明建议内容和理由）
        - 以下字段根据类型选择：
          * task类型需要: start_time, end_time, location, priority(1-3), category(personal/work/study/social)
          * trip类型需要: destination, start_date, end_date, budget(可选)

        时间格式：
        - start_time/end_time: ISO 8601格式，如 "2024-01-15T09:00:00"
        - start_date/end_date: ISO 8601格式，如 "2024-01-15"

        规划原则：
        1. 避免与用户现有任务冲突
        2. 合理分配时间，不要过于密集
        3. 优先安排高优先级事项
        4. 考虑日常作息规律（工作时间9:00-18:00，晚上20:00-22:00适合学习休息）
        5. 如果用户明确指定了月份（如"8月份"、"七月"等），旅行计划必须在该月份内安排

        示例输出：
        [
            {"type":"task","title":"团队周会","description":"建议每周一上午10点召开团队周会，便于总结上周工作和规划本周任务","start_time":"2024-01-15T10:00:00","end_time":"2024-01-15T11:00:00","location":"会议室A","priority":2,"category":"work"}
        ]

        如果无法生成有效建议，输出空数组[]
        """
        
        user_message = f"""用户需求：{prompt}
        
现有任务数据：
{existing_data}

请根据用户需求和现有任务安排，生成合理的任务/行程建议，严格按照JSON格式输出。
注意：不要生成冲突解决相关的任务，只生成实际的行程规划和有用的建议。"""
        
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f"Bearer {app.config['DEEPSEEK_API_KEY']}"
        }
        
        data = {
            'model': 'deepseek-chat',
            'messages': [
                {'role': 'system', 'content': system_prompt.strip()},
                {'role': 'user', 'content': user_message}
            ],
            'temperature': 0.5,
            'max_tokens': 1500
        }
        
        response = requests.post(app.config['DEEPSEEK_API_URL'], headers=headers, json=data, timeout=60)
        response.raise_for_status()
        
        result = response.json()['choices'][0]['message']['content']
        
        result = result.strip()
        if result.startswith('```'):
            result = result[3:]
        if result.endswith('```'):
            result = result[:-3]
        result = result.strip()
        
        try:
            suggestions = json.loads(result)
            if isinstance(suggestions, list):
                return validate_and_clean_suggestions(suggestions)
            else:
                return generate_ai_suggestions(prompt, existing_tasks, existing_trips)
        except json.JSONDecodeError:
            return generate_ai_suggestions(prompt, existing_tasks, existing_trips)
    
    except Exception as e:
        print(f"[ERROR] DeepSeek API call failed: {str(e)}")
        return generate_ai_suggestions(prompt, existing_tasks, existing_trips)

def validate_and_clean_suggestions(suggestions):
    cleaned = []
    today = datetime.now()
    
    for item in suggestions:
        if not isinstance(item, dict) or 'type' not in item or 'title' not in item:
            continue
        
        item_type = item['type']
        title = str(item['title'])[:50]
        description = str(item.get('description', ''))[:200]
        
        if item_type == 'task':
            raw_start_time = item.get('start_time', '')
            raw_end_time = item.get('end_time', '')
            
            # 验证start_time格式
            if raw_start_time and isinstance(raw_start_time, str) and len(raw_start_time) >= 16:
                try:
                    # 尝试解析时间
                    parsed_start = datetime.fromisoformat(raw_start_time.replace('Z', '+00:00'))
                    # 如果日期在过去，调整到下一年
                    if parsed_start < today:
                        parsed_start = parsed_start.replace(year=parsed_start.year + 1)
                    start_time = parsed_start.isoformat()
                    print(f"[DEBUG] Validated start_time: {raw_start_time} -> {start_time}")
                except Exception as e:
                    start_time = (today + timedelta(hours=1)).isoformat()
                    print(f"[DEBUG] Invalid start_time {raw_start_time}, using default: {start_time}")
            else:
                start_time = (today + timedelta(hours=1)).isoformat()
                print(f"[DEBUG] Empty or invalid start_time, using default: {start_time}")
            
            # 验证end_time格式
            if raw_end_time and isinstance(raw_end_time, str) and len(raw_end_time) >= 16:
                try:
                    parsed_end = datetime.fromisoformat(raw_end_time.replace('Z', '+00:00'))
                    if parsed_end < today:
                        parsed_end = parsed_end.replace(year=parsed_end.year + 1)
                    end_time = parsed_end.isoformat()
                    print(f"[DEBUG] Validated end_time: {raw_end_time} -> {end_time}")
                except Exception as e:
                    end_time = (today + timedelta(hours=2)).isoformat()
                    print(f"[DEBUG] Invalid end_time {raw_end_time}, using default: {end_time}")
            else:
                # 如果没有end_time，默认比start_time多1小时
                try:
                    start_dt = datetime.fromisoformat(start_time.replace('Z', '+00:00'))
                    end_time = (start_dt + timedelta(hours=1)).isoformat()
                except:
                    end_time = (today + timedelta(hours=2)).isoformat()
                print(f"[DEBUG] Empty end_time, calculated from start_time: {end_time}")
            
            cleaned.append({
                'type': 'task',
                'title': title,
                'description': description,
                'start_time': start_time,
                'end_time': end_time,
                'location': str(item.get('location', '')),
                'priority': max(1, min(3, int(item.get('priority', 2)))),
                'category': item.get('category', 'personal')
            })
        
        elif item_type == 'trip':
            start_date = item.get('start_date', (today + timedelta(days=7)).date().isoformat())
            end_date = item.get('end_date', (today + timedelta(days=10)).date().isoformat())
            
            cleaned.append({
                'type': 'trip',
                'title': title,
                'description': description,
                'destination': str(item.get('destination', '')),
                'start_date': start_date,
                'end_date': end_date,
                'budget': item.get('budget')
            })
    
    return cleaned

def generate_ai_suggestions(prompt, existing_tasks=None, existing_trips=None):
    """使用AI大模型生成智能建议"""
    suggestions = []
    today = datetime.now()
    
    existing_tasks = existing_tasks or []
    existing_trips = existing_trips or []
    
    existing_titles = set()
    for t in existing_tasks:
        existing_titles.add(t.get('title', '').lower())
    for t in existing_trips:
        existing_titles.add(t.get('title', '').lower())
    
    def call_ai_for_smart_suggestions(prompt, tasks, trips):
        """调用DeepSeek AI大模型生成智能建议"""
        if not DEEPSEEK_AVAILABLE:
            return None
        
        try:
            # 构建详细的日程信息
            tasks_info = []
            for t in tasks[:10]:  # 只取前10个任务
                try:
                    start = datetime.fromisoformat(t.get('start_time', '').replace('Z', '+00:00'))
                    end = datetime.fromisoformat(t.get('end_time', '').replace('Z', '+00:00'))
                    priority_map = {1: '高优先级', 2: '中优先级', 3: '低优先级'}
                    tasks_info.append({
                        'title': t.get('title', '未命名'),
                        'start': start.strftime('%Y年%m月%d日 %H:%M'),
                        'end': end.strftime('%Y年%m月%d日 %H:%M'),
                        'priority': priority_map.get(t.get('priority', 2), '中优先级'),
                        'category': t.get('category', '个人'),
                        'location': t.get('location', '未知地点')
                    })
                except:
                    pass
            
            trips_info = []
            for tr in trips[:5]:  # 只取前5个行程
                trips_info.append({
                    'title': tr.get('title', '未命名'),
                    'start': tr.get('start_date', ''),
                    'end': tr.get('end_date', ''),
                    'destination': tr.get('destination', '未知'),
                    'budget': tr.get('budget', 0)
                })
            
            # 构建智能prompt
            ai_prompt = f"""你是用户的智能日程助手。用户说："{prompt}"

用户当前日程安排：
任务列表（{len(tasks_info)}个）：
{json.dumps(tasks_info, ensure_ascii=False, indent=2)}

行程列表（{len(trips_info)}个）：
{json.dumps(trips_info, ensure_ascii=False, indent=2)}

请分析用户需求，生成智能建议。

## 重要要求：
1. 如果用户明确指定了月份（如"8月份"、"七月"等），旅行计划必须在该月份内安排
2. 如果用户指定了预算，旅行计划的预算预估要符合用户要求
3. 如果用户指定了行程天数，要按照指定天数规划
4. 旅行计划必须避开用户现有的任务和行程安排
5. 请仔细阅读用户的要求，确保所有约束条件都被满足

你的回复必须包含以下JSON格式（不要包含任何其他内容）：

{{
    "conflict_solutions": [
        {{
            "conflict_type": "时间冲突",
            "event1": "冲突事件1名称",
            "event2": "冲突事件2名称",
            "recommended_action": "建议执行哪个，应该考虑优先级、重要性、可调整性",
            "reason": "为什么建议这样做"
        }}
    ],
    "personalized_suggestions": [
        {{
            "type": "关怀建议",
            "title": "暖心建议标题",
            "description": "根据用户的工作量、日程密度给出的个性化关怀建议，比如'你最近任务很多，建议周末去吃顿火锅犒劳自己'",
            "action": "具体行动建议"
        }}
    ],
    "travel_plans": [
        {{
            "destination": "推荐目的地（根据用户偏好和现有行程选择）",
            "duration": 建议天数,
            "best_time": "最佳出发时间，格式为YYYY-MM-DD（严格按照用户指定的月份）",
            "reason": "为什么推荐这个目的地和时间",
            "attractions": ["景点1", "景点2"],
            "estimated_budget": 预算金额,
            "tips": "旅行注意事项"
        }}
    ],
    "smart_advice": [
        {{
            "category": "分类（工作/生活/健康）",
            "title": "建议标题",
            "description": "详细建议内容",
            "priority": "高/中/低"
        }}
    ]
}}

请确保：
1. 如果有日程冲突，给出明确的优先级判断和解决方案
2. 根据用户的日程密度和工作量，生成温暖的关怀建议
3. 如果用户提到旅行需求，严格按照用户指定的时间、预算、天数来制定旅行计划
4. 建议要有个性化，符合用户的实际情况
5. 只返回JSON，不要包含任何解释或其他文本"""
            
            headers = {
                'Content-Type': 'application/json',
                'Authorization': f"Bearer {Config.DEEPSEEK_API_KEY}"
            }
            
            data = {
                'model': 'deepseek-chat',
                'messages': [
                    {'role': 'user', 'content': ai_prompt}
                ],
                'temperature': 0.7,
                'max_tokens': 2000
            }
            
            response = requests.post(Config.DEEPSEEK_API_URL, headers=headers, json=data, timeout=60)
            response.raise_for_status()
            
            result = response.json()['choices'][0]['message']['content'].strip()
            
            # 提取JSON部分
            if '{' in result and '}' in result:
                json_start = result.find('{')
                json_end = result.rfind('}') + 1
                json_str = result[json_start:json_end]
                ai_result = json.loads(json_str)
                
                # 验证并修正AI返回的旅行计划
                for travel in ai_result.get('travel_plans', []):
                    best_time = travel.get('best_time', '')
                    if best_time:
                        # 检查日期格式
                        try:
                            travel_date = datetime.strptime(best_time, '%Y-%m-%d')
                            
                            # 检查是否符合用户指定的月份
                            month_match = re.search(r'([一二三四五六七八九十]|1[0-2])月', prompt)
                            if month_match:
                                month_name = month_match.group(1)
                                target_month = {'一':1, '二':2, '三':3, '四':4, '五':5, '六':6, 
                                            '七':7, '八':8, '九':9, '十':10, '11':11, '12':12}.get(month_name)
                                
                                if target_month and travel_date.month != target_month:
                                    # 修正日期到目标月份
                                    corrected_date = travel_date.replace(month=target_month)
                                    if corrected_date < today:
                                        corrected_date = corrected_date.replace(year=today.year + 1)
                                    travel['best_time'] = corrected_date.strftime('%Y-%m-%d')
                                    print(f"[DEBUG] 修正日期从 {best_time} 到 {travel['best_time']}")
                        except Exception as e:
                            print(f"[DEBUG] Error correcting date: {e}")
                
                return ai_result
            
        except Exception as e:
            print(f"[ERROR] DeepSeek AI call failed: {e}")
            return None
        
        return None
    
    def analyze_existing_tasks():
        analysis = {
            'has_conflicts': False,
            'busy_times': [],
            'free_times': [],
            'upcoming_events': [],
            'total_items': len(existing_tasks) + len(existing_trips)
        }
        
        all_events = []
        for t in existing_tasks:
            try:
                start = datetime.fromisoformat(t['start_time'].replace('Z', '+00:00'))
                end = datetime.fromisoformat(t['end_time'].replace('Z', '+00:00'))
                all_events.append({'start': start, 'end': end, 'title': t['title'], 'type': 'task'})
            except:
                pass
        
        for t in existing_trips:
            try:
                start = datetime.strptime(t['start_date'], '%Y-%m-%d')
                end = datetime.strptime(t['end_date'], '%Y-%m-%d')
                all_events.append({'start': start, 'end': end, 'title': t['title'], 'type': 'trip'})
            except:
                pass
        
        all_events.sort(key=lambda x: x['start'])
        
        for i in range(len(all_events) - 1):
            if all_events[i]['end'] > all_events[i+1]['start']:
                analysis['has_conflicts'] = True
                break
        
        analysis['upcoming_events'] = [e for e in all_events if e['start'] >= today and e['start'] <= today + timedelta(days=30)]
        
        return analysis
    
    # 调用AI生成智能建议
    ai_result = call_ai_for_smart_suggestions(prompt, existing_tasks, existing_trips)
    
    task_analysis = analyze_existing_tasks()
    
    # 如果AI调用成功，使用AI生成的结果
    if ai_result:
        # 处理冲突解决方案
        for conflict in ai_result.get('conflict_solutions', []):
            suggestions.append({
                'type': 'task',
                'title': f"📋 冲突解决：{conflict.get('event1', '任务1')} vs {conflict.get('event2', '任务2')}",
                'description': f"【问题】{conflict.get('conflict_type', '时间冲突')}\n\n【建议】{conflict.get('recommended_action', '待定')}\n\n【原因】{conflict.get('reason', '请根据实际情况判断')}",
                'start_time': today.replace(hour=10, minute=0).isoformat(),
                'end_time': today.replace(hour=11, minute=0).isoformat(),
                'priority': 1,
                'category': 'work'
            })
        
        # 处理个性化关怀建议
        for care in ai_result.get('personalized_suggestions', []):
            suggestions.append({
                'type': 'task',
                'title': care.get('title', '暖心建议'),
                'description': care.get('description', ''),
                'start_time': today.replace(hour=19, minute=0).isoformat(),
                'end_time': today.replace(hour=20, minute=0).isoformat(),
                'priority': 3,
                'category': 'personal',
                'is_care': True
            })
        
        # 处理旅行计划
        for travel in ai_result.get('travel_plans', []):
            try:
                start_date_str = travel.get('best_time', today.strftime('%Y-%m-%d'))
                start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
                duration = int(travel.get('duration', 3))
                
                # 检查并修正日期到用户指定的月份
                month_match = re.search(r'([一二三四五六七八九十]|1[0-2])月', prompt)
                if month_match:
                    month_name = month_match.group(1)
                    target_month = {'一':1, '二':2, '三':3, '四':4, '五':5, '六':6, 
                                '七':7, '八':8, '九':9, '十':10, '11':11, '12':12}.get(month_name)
                    
                    if target_month and start_date.month != target_month:
                        # 修正日期到目标月份的第一天
                        start_date = start_date.replace(month=target_month, day=1)
                        if start_date < today:
                            start_date = start_date.replace(year=today.year + 1)
                        print(f"[DEBUG] 日期修正: {start_date_str} -> {start_date.strftime('%Y-%m-%d')}")
                
                end_date = start_date + timedelta(days=duration - 1)
                
                attractions_str = '\n'.join([f"• {att}" for att in travel.get('attractions', [])])
                
                suggestions.append({
                    'type': 'trip',
                    'title': f"✈️ {travel.get('destination', '旅行计划')}",
                    'description': f"【目的地】{travel.get('destination', '待定')}\n\n【推荐理由】{travel.get('reason', '风景优美，适合放松')}\n\n【推荐景点】\n{attractions_str}\n\n【预算预估】约¥{travel.get('estimated_budget', 3000)}\n\n【注意事项】{travel.get('tips', '提前预订，注意安全')}",
                    'destination': travel.get('destination', '待定'),
                    'start_date': start_date.strftime('%Y-%m-%d'),
                    'end_date': end_date.strftime('%Y-%m-%d'),
                    'budget': travel.get('estimated_budget', 3000),
                    'priority': 2
                })
            except Exception as e:
                print(f"[ERROR] Failed to parse travel plan: {e}")
        
        # 处理智能建议
        for advice in ai_result.get('smart_advice', []):
            priority_map = {'高': 1, '中': 2, '低': 3}
            suggestions.append({
                'type': 'task',
                'title': f"💡 {advice.get('title', '智能建议')}",
                'description': advice.get('description', ''),
                'start_time': (today + timedelta(days=1)).replace(hour=14, minute=0).isoformat(),
                'end_time': (today + timedelta(days=1)).replace(hour=15, minute=0).isoformat(),
                'priority': priority_map.get(advice.get('priority', '中'), 2),
                'category': 'personal'
            })
        
        # 全局日期修正：确保所有旅行计划的日期符合用户指定的月份
        month_match = re.search(r'([一二三四五六七八九十]|1[0-2])月', prompt)
        if month_match:
            month_name = month_match.group(1)
            target_month = {'一':1, '二':2, '三':3, '四':4, '五':5, '六':6, 
                        '七':7, '八':8, '九':9, '十':10, '11':11, '12':12}.get(month_name)
            
            if target_month:
                for s in suggestions:
                    if s.get('type') == 'trip' and 'start_date' in s:
                        try:
                            start_date = datetime.strptime(s['start_date'], '%Y-%m-%d')
                            if start_date.month != target_month:
                                corrected_date = start_date.replace(month=target_month, day=1)
                                if corrected_date < today:
                                    corrected_date = corrected_date.replace(year=today.year + 1)
                                s['start_date'] = corrected_date.strftime('%Y-%m-%d')
                                end_date = corrected_date + timedelta(days=int(s.get('budget', 3000)/500) - 1)
                                s['end_date'] = end_date.strftime('%Y-%m-%d')
                                print(f"[DEBUG] 全局日期修正: {s['title']} -> {s['start_date']}")
                        except Exception as e:
                            print(f"[DEBUG] 全局日期修正失败: {e}")
        
        return suggestions
    
    # 如果AI调用失败，使用备用规则生成建议
    print("[WARNING] AI call failed, using fallback rules")
    
    def parse_user_constraints(prompt_text):
        """解析用户约束条件"""
        constraints = {
            'avoid_weekends': bool(re.search(r'(避开周末|周末除外|不要周末|非周末)', prompt_text)),
            'avoid_holidays': bool(re.search(r'(避开节假日|节假日除外|不要节假日)', prompt_text)),
            'prefer_workdays': bool(re.search(r'(工作日|上班时间|周一到周五)', prompt_text)),
            'duration': extract_duration(prompt_text),
            'destination': extract_destination(prompt_text),
            'preferred_month': extract_month(prompt_text),
            'keywords': extract_keywords(prompt_text)
        }
        return constraints
    
    def extract_duration(text):
        """提取持续时间"""
        duration_patterns = [
            (r'(\d+)天', lambda m: int(m.group(1))),
            (r'(\d+)个?星期', lambda m: int(m.group(1)) * 7),
            (r'(\d+)周', lambda m: int(m.group(1)) * 7),
            (r'(\d+)小时', lambda m: int(m.group(1)) / 24)
        ]
        for pattern, converter in duration_patterns:
            match = re.search(pattern, text)
            if match:
                return converter(match)
        return 3  # 默认3天
    
    def extract_destination(text):
        """提取目的地"""
        dest_patterns = [
            r'去\s*([\u4e00-\u9fa5]+?)旅游',
            r'去\s*([\u4e00-\u9fa5]+?)旅行',
            r'到\s*([\u4e00-\u9fa5]+?)旅游',
            r'到\s*([\u4e00-\u9fa5]+?)旅行',
            r'旅游\s*([\u4e00-\u9fa5]+)',
            r'旅行\s*([\u4e00-\u9fa5]+)',
            r'出差\s*([\u4e00-\u9fa5]+)'
        ]
        for pattern in dest_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()
        return '待确定'
    
    def extract_month(text):
        """提取月份"""
        month_match = re.search(r'([一二三四五六七八九十]|1[0-2])月', text)
        if month_match:
            month_name = month_match.group(1)
            return {'一':1, '二':2, '三':3, '四':4, '五':5, '六':6, 
                    '七':7, '八':8, '九':9, '十':10, '11':11, '12':12}.get(month_name)
        return None
    
    def extract_keywords(text):
        """提取关键词"""
        keywords = []
        if re.search(r'(旅行|旅游|出差|trip)', text, re.IGNORECASE):
            keywords.append('travel')
        if re.search(r'(会议|meeting|开会)', text, re.IGNORECASE):
            keywords.append('meeting')
        if re.search(r'(规划|安排|整理)', text, re.IGNORECASE):
            keywords.append('plan')
        if re.search(r'(考试|答辩|作业|学习)', text, re.IGNORECASE):
            keywords.append('study')
        return keywords
    
    def find_best_travel_dates(duration_days=3, avoid_weekends=True, preferred_month=None):
        """找到最佳旅行日期，避开周末和现有安排"""
        start_date = today + timedelta(days=1)
        end_date = today + timedelta(days=90)
        
        best_options = []
        current_date = start_date
        
        print(f"[DEBUG] Looking for {duration_days}-day travel, avoid weekends: {avoid_weekends}, preferred month: {preferred_month}")
        print(f"[DEBUG] Existing trips count: {len(existing_trips)}")
        print(f"[DEBUG] Existing tasks count: {len(existing_tasks)}")
        
        while current_date <= end_date - timedelta(days=duration_days):
            current_date_only = current_date.date()
            
            if preferred_month and current_date_only.month != preferred_month:
                current_date += timedelta(days=1)
                continue
            
            weekend_conflict = False
            for i in range(int(duration_days)):
                check_date = current_date_only + timedelta(days=i)
                if check_date.weekday() >= 5:
                    weekend_conflict = True
                    break
            
            if avoid_weekends and weekend_conflict:
                current_date += timedelta(days=1)
                continue
            
            trip_end = current_date_only + timedelta(days=duration_days - 1)
            conflict = False
            conflicting_events = []
            
            for event in existing_trips:
                try:
                    event_start = datetime.strptime(event['start_date'], '%Y-%m-%d').date()
                    event_end = datetime.strptime(event['end_date'], '%Y-%m-%d').date()
                    
                    if not (trip_end < event_start or current_date_only > event_end):
                        conflict = True
                        conflicting_events.append(f"行程: {event.get('title', '未知')} ({event_start} ~ {event_end})")
                        break
                except Exception as e:
                    print(f"[DEBUG] Error checking trip conflict: {e}")
            
            if not conflict:
                for event in existing_tasks:
                    try:
                        event_start = datetime.fromisoformat(event['start_time'].replace('Z', '+00:00')).date()
                        event_end = datetime.fromisoformat(event['end_time'].replace('Z', '+00:00')).date()
                        
                        if not (trip_end < event_start or current_date_only > event_end):
                            conflict = True
                            conflicting_events.append(f"任务: {event.get('title', '未知')} ({event_start} ~ {event_end})")
                            break
                    except Exception as e:
                        print(f"[DEBUG] Error checking task conflict: {e}")
            
            if conflicting_events:
                print(f"[DEBUG] Conflict on {current_date_only}: {', '.join(conflicting_events)}")
            
            if not conflict:
                print(f"[DEBUG] Found good date: {current_date_only} ~ {trip_end}")
                best_options.append({
                    'start': current_date_only,
                    'end': trip_end
                })
            
            current_date += timedelta(days=1)
        
        print(f"[DEBUG] Found {len(best_options)} good options")
        return best_options[:5]
    
    def analyze_existing_tasks():
        analysis = {
            'has_conflicts': False,
            'conflicts': [],
            'busy_times': [],
            'free_times': [],
            'upcoming_events': [],
            'total_items': len(existing_tasks) + len(existing_trips),
            'priority_summary': {'high': 0, 'medium': 0, 'low': 0},
            'recommendations': []
        }
        
        all_events = []
        for t in existing_tasks:
            try:
                start = datetime.fromisoformat(t['start_time'].replace('Z', '+00:00'))
                end = datetime.fromisoformat(t['end_time'].replace('Z', '+00:00'))
                all_events.append({'start': start, 'end': end, 'title': t['title'], 'type': 'task', 'priority': t.get('priority', 2)})
                if t.get('priority') == 1:
                    analysis['priority_summary']['high'] += 1
                elif t.get('priority') == 2:
                    analysis['priority_summary']['medium'] += 1
                else:
                    analysis['priority_summary']['low'] += 1
            except:
                pass
        
        for t in existing_trips:
            try:
                start = datetime.strptime(t['start_date'], '%Y-%m-%d')
                end = datetime.strptime(t['end_date'], '%Y-%m-%d')
                all_events.append({'start': start, 'end': end, 'title': t['title'], 'type': 'trip', 'priority': 2})
                analysis['priority_summary']['medium'] += 1
            except:
                pass
        
        all_events.sort(key=lambda x: x['start'])
        
        for i in range(len(all_events) - 1):
            if all_events[i]['end'] > all_events[i+1]['start']:
                analysis['has_conflicts'] = True
                analysis['conflicts'].append({
                    'event1': all_events[i],
                    'event2': all_events[i+1]
                })
        
        analysis['upcoming_events'] = [e for e in all_events if e['start'] >= today and e['start'] <= today + timedelta(days=30)]
        
        return analysis
    
    task_analysis = analyze_existing_tasks()
    constraints = parse_user_constraints(prompt)
    
    month_num = constraints['preferred_month']
    if month_num:
        month_str = ['一月','二月','三月','四月','五月','六月','七月','八月','九月','十月','十一月','十二月'][month_num-1]
        if task_analysis['upcoming_events']:
            event_list = "\n".join([f"• {e['title']}" for e in task_analysis['upcoming_events'][:5]])
            suggestions.append({
                'type': 'task',
                'title': f'{month_str}任务回顾',
                'description': f'您{month_str}的已有安排包括：\n{event_list}',
                'start_time': today.replace(hour=10, minute=0).isoformat(),
                'end_time': today.replace(hour=11, minute=0).isoformat(),
                'location': '办公室',
                'priority': 2,
                'category': 'personal'
            })
        else:
            suggestions.append({
                'type': 'task',
                'title': f'规划{month_str}安排',
                'description': f'您{month_str}目前没有安排，建议规划本月的重要事项',
                    'start_time': (today + timedelta(days=3)).replace(hour=18, minute=0).isoformat(),
                    'end_time': (today + timedelta(days=3)).replace(hour=19, minute=0).isoformat(),
                    'priority': 2,
                    'category': 'personal'
                })
    
    if re.search(r'(会议|meeting|开会|讨论)', prompt, re.IGNORECASE):
        if '团队周会' not in existing_titles:
            suggestions.append({
                'type': 'task',
                'title': '团队周会',
                'description': '建议每周安排团队会议，同步工作进展',
                'start_time': (today + timedelta(days=(7 - today.weekday()) % 7)).replace(hour=10, minute=0).isoformat(),
                'end_time': (today + timedelta(days=(7 - today.weekday()) % 7)).replace(hour=11, minute=0).isoformat(),
                'location': '会议室',
                'priority': 2,
                'category': 'work'
            })
    
    def get_travel_destination_info(destination):
        """获取旅游目的地信息"""
        destinations = {
            '云南大理': {
                'description': '大理是云南著名的旅游城市，以苍山洱海、古城风情闻名',
                'best_months': ['3月', '4月', '5月', '9月', '10月'],
                'weather': '七月平均气温15-25°C，凉爽舒适',
                'attractions': ['洱海环湖', '大理古城', '苍山索道', '喜洲古镇', '双廊古镇'],
                'transport': '建议自驾或包车，方便环湖游',
                'accommodation': ['大理古城周边民宿', '洱海海景客栈', '双廊海景酒店'],
                'budget_per_day': 500,
                'notes': ['注意防晒', '高原紫外线强', '提前预订住宿']
            },
            '肇庆': {
                'description': '肇庆是广东历史文化名城，七星岩、鼎湖山是著名景点',
                'best_months': ['4月', '5月', '9月', '10月'],
                'weather': '七月炎热潮湿，注意防暑降温',
                'attractions': ['七星岩', '鼎湖山', '肇庆古城墙', '德庆盘龙峡'],
                'transport': '高铁方便，市内打车或公交',
                'accommodation': ['七星岩附近酒店', '肇庆市区酒店'],
                'budget_per_day': 400,
                'notes': ['夏季多雨', '注意防蚊虫']
            },
            '桂林': {
                'description': '桂林山水甲天下，漓江风光美不胜收',
                'best_months': ['4月', '5月', '9月', '10月'],
                'weather': '七月炎热，雨水较多',
                'attractions': ['漓江游船', '阳朔西街', '遇龙河漂流', '兴坪古镇'],
                'transport': '建议包车或跟团',
                'accommodation': ['阳朔民宿', '桂林市区酒店'],
                'budget_per_day': 600,
                'notes': ['雨季注意安全', '提前订漓江船票']
            },
            '丽江': {
                'description': '丽江古城世界文化遗产，玉龙雪山壮丽非凡',
                'best_months': ['4月', '5月', '9月', '10月'],
                'weather': '七月凉爽，平均气温15-24°C',
                'attractions': ['丽江古城', '玉龙雪山', '束河古镇', '泸沽湖'],
                'transport': '建议自驾或包车',
                'accommodation': ['丽江古城客栈', '束河古镇民宿'],
                'budget_per_day': 550,
                'notes': ['高原反应注意', '古城商业化较浓']
            },
            '厦门': {
                'description': '海滨城市，鼓浪屿风情独特',
                'best_months': ['10月', '11月', '3月', '4月'],
                'weather': '七月炎热潮湿',
                'attractions': ['鼓浪屿', '厦门大学', '曾厝垵', '环岛路'],
                'transport': '市内公交便捷',
                'accommodation': ['鼓浪屿民宿', '曾厝垵客栈'],
                'budget_per_day': 450,
                'notes': ['暑期游客多', '提前预订鼓浪屿住宿']
            },
            '张家界': {
                'description': '世界自然遗产，奇峰异石令人惊叹',
                'best_months': ['4月', '5月', '9月', '10月'],
                'weather': '七月较热，注意防暑',
                'attractions': ['天门山', '张家界国家森林公园', '玻璃栈道', '黄龙洞'],
                'transport': '景区内环保车',
                'accommodation': ['张家界市区酒店', '景区门口客栈'],
                'budget_per_day': 500,
                'notes': ['山区温差大', '穿舒适运动鞋']
            }
        }
        return destinations.get(destination, {
            'description': '待确定目的地',
            'best_months': ['全年'],
            'weather': '请查询当地天气',
            'attractions': [],
            'transport': '待定',
            'accommodation': [],
            'budget_per_day': 400,
            'notes': []
        })

    def generate_detailed_travel_plan(destination, duration, start_date):
        """生成详细的旅行计划"""
        info = get_travel_destination_info(destination)
        plan = []
        
        day1 = start_date
        for day in range(duration):
            current_date = day1 + timedelta(days=day)
            plan.append({
                'day': day + 1,
                'date': current_date.strftime('%Y年%m月%d日'),
                'day_of_week': ['周一', '周二', '周三', '周四', '周五', '周六', '周日'][current_date.weekday()],
                'suggestion': f'Day{day+1}: ' + get_day_suggestion(destination, day, duration)
            })
        
        return {
            'destination': destination,
            'duration': duration,
            'start_date': start_date.strftime('%Y-%m-%d'),
            'end_date': (start_date + timedelta(days=duration-1)).strftime('%Y-%m-%d'),
            'destination_info': info,
            'daily_plan': plan,
            'total_budget': info['budget_per_day'] * duration,
            'transport': info['transport'],
            'accommodation': info['accommodation'],
            'notes': info['notes']
        }
    
    def get_day_suggestion(destination, day, total_days):
        """获取每天的行程建议"""
        suggestions = {
            '云南大理': [
                '抵达大理，入住酒店，下午逛大理古城，品尝当地美食',
                '洱海环湖一日游，可租电动车或包车',
                '苍山索道，俯瞰大理全景',
                '喜洲古镇体验白族文化',
                '双廊古镇休闲，欣赏洱海日落',
                '返程，购买特产'
            ],
            '肇庆': [
                '抵达肇庆，入住酒店，七星岩半日游',
                '鼎湖山一日游，感受天然氧吧',
                '肇庆古城墙、披云楼历史游',
                '德庆盘龙峡或竹海大观',
                '返程'
            ],
            '桂林': [
                '抵达桂林，象鼻山公园',
                '漓江游船至阳朔',
                '阳朔西街夜游',
                '遇龙河漂流，骑行遇龙河',
                '兴坪古镇，20元人民币背景',
                '返程'
            ],
            '丽江': [
                '抵达丽江，入住古城客栈',
                '丽江古城深度游',
                '玉龙雪山一日游',
                '束河古镇休闲',
                '泸沽湖一日游',
                '返程'
            ],
            '厦门': [
                '抵达厦门，入住酒店，环岛路骑行',
                '鼓浪屿一日游',
                '厦门大学，曾厝垵',
                '南普陀寺，沙坡尾艺术区',
                '返程'
            ],
            '张家界': [
                '抵达张家界，入住酒店',
                '张家界国家森林公园一日',
                '天门山景区',
                '玻璃栈道，黄龙洞',
                '返程'
            ]
        }
        
        dest_suggestions = suggestions.get(destination, [])
        if day < len(dest_suggestions):
            return dest_suggestions[day]
        return f'自由活动或根据兴趣安排景点游览'
    
    if 'travel' in constraints['keywords']:
        dest = constraints['destination']
        duration = int(constraints['duration'])
        best_dates = find_best_travel_dates(
            duration_days=duration,
            avoid_weekends=constraints['avoid_weekends'],
            preferred_month=constraints['preferred_month']
        )
        
        if best_dates:
            best_option = best_dates[0]
            trip_title = f'{dest}旅行计划' if dest != '待确定' else '旅行计划'
            
            travel_plan = generate_detailed_travel_plan(dest, duration, best_option['start'])
            
            suggestions.append({
                'type': 'trip',
                'title': trip_title,
                'description': f'建议于{best_option["start"].strftime("%Y年%m月%d日")}出发前往{dest}，行程{duration}天\n\n【目的地简介】\n{travel_plan["destination_info"]["description"]}\n\n【天气情况】\n{travel_plan["destination_info"]["weather"]}\n\n【推荐景点】\n' + '\n'.join([f'• {attraction}' for attraction in travel_plan["destination_info"]["attractions"]]) + f'\n\n【交通方式】\n{travel_plan["transport"]}\n\n【住宿推荐】\n' + '\n'.join([f'• {hotel}' for hotel in travel_plan["destination_info"]["accommodation"]]) + f'\n\n【预算预估】\n总计约¥{travel_plan["total_budget"]}（约¥{travel_plan["destination_info"]["budget_per_day"]}/天）\n\n【注意事项】\n' + '\n'.join([f'• {note}' for note in travel_plan["destination_info"]["notes"]]),
                'destination': dest,
                'start_date': best_option['start'].isoformat(),
                'end_date': best_option['end'].isoformat(),
                'budget': travel_plan['total_budget'],
                'priority': 2,
                'travel_plan': travel_plan
            })
        else:
            suggestions.append({
                'type': 'task',
                'title': '规划旅行时间',
                'description': '根据您的约束条件和现有安排，近期暂无合适的旅行时间。建议调整出行时间、减少行程天数，或等现有行程结束后再安排。',
                'start_time': today.replace(hour=15, minute=0).isoformat(),
                'end_time': today.replace(hour=16, minute=0).isoformat(),
                'priority': 2,
                'category': 'personal'
            })

    if 'meeting' in constraints['keywords']:
        if '团队周会' not in existing_titles:
            suggestions.append({
                'type': 'task',
                'title': '团队周会',
                'description': '建议每周安排团队会议，同步工作进展',
                'start_time': (today + timedelta(days=(7 - today.weekday()) % 7)).replace(hour=10, minute=0).isoformat(),
                'end_time': (today + timedelta(days=(7 - today.weekday()) % 7)).replace(hour=11, minute=0).isoformat(),
                'location': '会议室',
                'priority': 2,
                'category': 'work'
            })

    if 'plan' in constraints['keywords']:
        if '任务规划' not in existing_titles:
            suggestions.append({
                'type': 'task',
                'title': '任务规划',
                'description': '建议定期进行任务规划，合理安排时间',
                'start_time': today.replace(hour=9, minute=0).isoformat(),
                'end_time': today.replace(hour=10, minute=0).isoformat(),
                'location': '办公室',
                'priority': 2,
                'category': 'personal'
            })

    if task_analysis['has_conflicts']:
        if '处理任务冲突' not in existing_titles:
            suggestions.append({
                'type': 'task',
                'title': '处理任务冲突',
                'description': '检测到您的任务存在冲突，建议及时调整',
                'start_time': today.replace(hour=12, minute=0).isoformat(),
                'end_time': today.replace(hour=13, minute=0).isoformat(),
                'priority': 1,
                'category': 'work'
            })

    if task_analysis['total_items'] == 0:
        suggestions.append({
            'type': 'task',
            'title': '首次安排',
            'description': '您还没有任何任务安排，建议添加第一个任务',
            'start_time': (today + timedelta(days=1)).replace(hour=10, minute=0).isoformat(),
            'end_time': (today + timedelta(days=1)).replace(hour=11, minute=0).isoformat(),
            'location': '办公室',
            'priority': 2,
            'category': 'work'
        })

    return suggestions

@app.route('/api/upload/file', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        user_id = request.form.get('user_id')
        if not user_id:
            return jsonify({'error': 'User ID is required'}), 400
        
        file_ext = file.filename.split('.')[-1].lower()
        
        supported_extensions = ['txt', 'md', 'csv', 'docx', 'doc', 'xlsx', 'xls', 'pdf', 'json', 'xml', 'html']
        if file_ext not in supported_extensions:
            return jsonify({'error': f'不支持的文件类型: {file_ext}。支持的类型: {", ".join(supported_extensions)}'}), 400
        
        content = extract_file_content(file, file_ext)
        
        if not content:
            return jsonify({'error': '无法提取文件内容'}), 400
        
        parsed_items = parse_text_content_with_ai(content, user_id)
        
        return jsonify({
            'message': 'File uploaded and parsed successfully',
            'filename': file.filename,
            'file_type': file_ext,
            'items': parsed_items
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def extract_file_content(file, file_ext):
    """根据文件类型提取内容"""
    try:
        if file_ext in ['txt', 'md', 'csv', 'json', 'xml', 'html']:
            return file.read().decode('utf-8', errors='ignore')
        
        elif file_ext == 'docx':
            try:
                from docx import Document
                doc = Document(file)
                content = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        content.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        if any(row_text):
                            content.append(' | '.join(row_text))
                return '\n'.join(content)
            except ImportError:
                print("[WARNING] python-docx not installed, trying basic extraction")
                return file.read().decode('utf-8', errors='ignore')
        
        elif file_ext == 'doc':
            try:
                file.seek(0)
                file_data = file.read()
                
                if file_data[:8] == b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1':
                    print("[DEBUG] Detected OLE2 .doc file")
                    try:
                        import win32com.client
                        import pythoncom
                        pythoncom.CoInitialize()
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        
                        import tempfile
                        import os
                        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                            temp_file.write(file_data)
                            temp_path = temp_file.name
                        
                        try:
                            doc = word.Documents.Open(temp_path)
                            content = doc.Content.Text
                            doc.Close()
                            word.Quit()
                            pythoncom.CoUninitialize()
                            os.unlink(temp_path)
                            return content
                        except Exception as e:
                            print(f"[WARNING] Win32COM failed: {str(e)}")
                            word.Quit()
                            pythoncom.CoUninitialize()
                            os.unlink(temp_path)
                    except ImportError:
                        print("[WARNING] pywin32 not installed")
                
                try:
                    from docx import Document
                    file.seek(0)
                    doc = Document(file)
                    content = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            content.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells]
                            if any(row_text):
                                content.append(' | '.join(row_text))
                    return '\n'.join(content)
                except Exception as e:
                    print(f"[WARNING] python-docx failed: {str(e)}")
                
                try:
                    import subprocess
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                        temp_file.write(file_data)
                        temp_path = temp_file.name
                    
                    try:
                        result = subprocess.run(
                            ['antiword', temp_path],
                            capture_output=True,
                            text=True,
                            timeout=30
                        )
                        if result.returncode == 0:
                            return result.stdout
                        else:
                            print(f"[WARNING] antiword failed, trying catdoc")
                            result = subprocess.run(
                                ['catdoc', temp_path],
                                capture_output=True,
                                text=True,
                                timeout=30
                            )
                            if result.returncode == 0:
                                return result.stdout
                            else:
                                print(f"[WARNING] catdoc also failed")
                    finally:
                        import os
                        os.unlink(temp_path)
                        
                    return file_data.decode('latin-1', errors='ignore')
                except Exception as e:
                    print(f"[WARNING] Failed to extract .doc file: {str(e)}")
                    return file_data.decode('latin-1', errors='ignore')
            except Exception as e:
                print(f"[ERROR] .doc extraction error: {str(e)}")
                return None
        
        elif file_ext in ['xlsx', 'xls']:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file)
                content = []
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    content.append(f"=== {sheet_name} ===")
                    for row in sheet.iter_rows(values_only=True):
                        row_text = [str(cell) if cell else '' for cell in row]
                        if any(row_text):
                            content.append(' | '.join(row_text))
                return '\n'.join(content)
            except ImportError:
                print("[WARNING] openpyxl not installed, trying basic extraction")
                return file.read().decode('utf-8', errors='ignore')
        
        elif file_ext == 'pdf':
            try:
                import pdfplumber
                with pdfplumber.open(file) as pdf:
                    content = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            content.append(text)
                    return '\n'.join(content)
            except ImportError:
                print("[WARNING] pdfplumber not installed")
                return None
        
        else:
            return file.read().decode('utf-8', errors='ignore')
    
    except Exception as e:
        print(f"[ERROR] Failed to extract content: {str(e)}")
        return None

def parse_text_content_with_ai(content, user_id):
    print(f"[AI PARSE] Parsing content with AI: {content[:100]}...")
    
    if DEEPSEEK_AVAILABLE:
        try:
            return ai_parse_text_content(content, user_id)
        except Exception as e:
            print(f"[AI PARSE] AI parsing failed: {str(e)[:100]}")
            return parse_text_content(content, user_id)
    else:
        return parse_text_content(content, user_id)

def ai_parse_text_content(content, user_id):
    try:
        current_year = datetime.now().year
        system_prompt = f"""
        你是一个专业的任务信息提取助手，擅长从文本中智能识别和提取任务、行程相关信息。
        
        核心能力：
        1. 从文本中提取日期（如：6月20日、{current_year}-06-20、明天下午3点等）
        2. 提取时间（开始时间、结束时间）
        3. 提取地点（会议室、办公室、地址等）
        4. 提取事件标题和描述
        5. 智能分类为任务/行程
        
        类型判断规则：
        - task（任务）：会议、培训、考试、面试、开会、通知、会议记录等日常工作或事务
        - trip（行程）：出差、旅行、旅游、景点游玩、度假等需要出行的活动
        
        时间处理规则：
        - 对于task（任务）：如果提到"大概一个小时"、"一个小时"等时间长度描述，结束时间=开始时间+对应时长
        - 如果文本中没有明确年份，默认使用当前年份{current_year}年
        - 如果解析出的日期已经过去（早于今天），请自动调整到{current_year+1}年
        - 对于task（任务）：如果没有明确时间长度，默认结束时间=开始时间+1小时
        - 对于trip（行程）：如果提到持续天数，设置合理的结束日期；否则默认3天

        输出格式要求：
        请输出一个JSON数组，每个元素包含以下字段：
        - type: 类型，可选值: "task"（任务）, "trip"（行程）
        - title: 标题（简洁明了，不超过20字）
        - description: 描述（详细说明）
        - start_time: 开始时间，ISO 8601格式，如 "{current_year}-06-20T14:00:00"
        - end_time: 结束时间（任务必填，格式同start_time）
        - location: 地点（可选）
        - priority: 优先级1-3（1最高）
        - category: 分类（personal/work/study/social）

        如果无法识别到有效信息，输出空数组[]
        """
        
        user_message = f"""请从以下文本中提取任务或行程信息：

{content}

请按照JSON格式输出提取的信息。"""
        
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f"Bearer {app.config['DEEPSEEK_API_KEY']}"
        }
        
        data = {
            'model': 'deepseek-chat',
            'messages': [
                {'role': 'system', 'content': system_prompt.strip()},
                {'role': 'user', 'content': user_message}
            ],
            'temperature': 0.3,
            'max_tokens': 1000
        }
        
        response = requests.post(app.config['DEEPSEEK_API_URL'], headers=headers, json=data, timeout=60)
        response.raise_for_status()
        
        result = response.json()['choices'][0]['message']['content'].strip()
        if result.startswith('```'):
            result = result[3:]
        if result.endswith('```'):
            result = result[:-3]
        result = result.strip()
        
        try:
            suggestions = json.loads(result)
            if isinstance(suggestions, list):
                return validate_and_clean_suggestions(suggestions)
            else:
                return parse_text_content(content, user_id)
        except json.JSONDecodeError:
            return parse_text_content(content, user_id)
            
    except Exception as e:
        print(f"[AI PARSE] DeepSeek API error: {str(e)}")
        return parse_text_content(content, user_id)

def parse_text_content(content, user_id):
    items = []
    
    print(f"[DEBUG] Parsing content length: {len(content)} chars")
    print(f"[DEBUG] Content: {repr(content)}")
    
    # 智能分类关键词
    trip_keywords = ['出差', '旅行', '旅游', '行程', '出发', '抵达', '航班', '机票', '火车', '高铁', 
                     '酒店', '住宿', '景点', '游玩', '度假', '目的地', '行程安排', '旅游团', 
                     '签证', '护照', '行李', '接送', '包车', '自驾', '公路', '机场', '车站',
                     '游玩', '观光', '游览', '返程', '行程表', '路线']
    
    task_keywords = ['任务', '待办', '完成', '提交', '审核', '审批', '处理', '跟进', '回复',
                     '整理', '撰写', '编写', '修改', '更新', '确认', '检查', '测试', '部署',
                     '截止', 'deadline', '紧急', '重要', '优先', '待处理', '进行中', '已完成',
                     '会议', '培训', '考试', '面试', '预约', '活动', '讲座', '课程', 
                     '报告', '讨论', '开会', '座谈', '研讨会', '座谈会', '晨会', '晚会',
                     '聚餐', '派对', '庆典', '仪式', '签约', '发布会', '展览', '展会',
                     '通知', '会议记录', '带上', '做好', '监考']
    
    today = datetime.now()
    year = today.year
    
    rows = content.split('\n')
    table_rows = []
    header_found = False
    
    for row in rows:
        if '|' in row:
            parts = [p.strip() for p in row.split('|')]
            if len(parts) >= 2:
                if not header_found and ('时间' in parts or '日期' in parts or '日程' in parts):
                    header_found = True
                elif header_found or any(re.search(r'(\d{1,2})月(\d{1,2})日', p) for p in parts):
                    table_rows.append(parts)
    
    if len(table_rows) > 0:
        print(f"[DEBUG] Found {len(table_rows)} table rows")
        for row_parts in table_rows:
            if len(row_parts) < 2:
                continue
            
            row_content = ' | '.join(row_parts)
            if not row_content.strip():
                continue
            
            # 先尝试匹配完整格式 "2026年6月22日"
            full_date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', row_content)
            # 再尝试匹配省略年份格式 "6月22日"
            short_date_match = re.search(r'(\d{1,2})月(\d{1,2})日', row_content)
            # 支持带空格的时间格式，如 "9: 00-11: 00" 或 "15:00-17:00"
            # 同时支持全角冒号和半角冒号
            time_range_match = re.search(r'(\d{1,2})[:：]\s*(\d{2})\s*[-~至]\s*(\d{1,2})[:：]\s*(\d{2})', row_content)
            
            if full_date_match:
                year = int(full_date_match.group(1))
                month = int(full_date_match.group(2))
                day = int(full_date_match.group(3))
                print(f"[DEBUG] Found full date: {year}年{month}月{day}日")
            elif short_date_match:
                month = int(short_date_match.group(1))
                day = int(short_date_match.group(2))
                print(f"[DEBUG] Found short date: {month}月{day}日, using year {year}")
            else:
                print(f"[DEBUG] No date found in row: {row_content}")
                continue
            
            if month < 1 or month > 12:
                month = today.month
            if day < 1 or day > 31:
                day = today.day
            
            parsed_date = f"{year:04d}-{month:02d}-{day:02d}"
            
            # 检查日期是否在过去，如果是，自动调整到下一年
            date_obj = datetime(year, month, day)
            if date_obj < today:
                # 如果日期已经过去，尝试下一年
                next_year = year + 1
                parsed_date = f"{next_year:04d}-{month:02d}-{day:02d}"
                print(f"[DEBUG] Date {year}-{month}-{day} is in the past, adjusted to {next_year}-{month}-{day}")
            
            start_time = "09:00:00"
            end_time = "10:00:00"
            
            if time_range_match:
                start_hour = int(time_range_match.group(1))
                start_minute = int(time_range_match.group(2))
                end_hour = int(time_range_match.group(3))
                end_minute = int(time_range_match.group(4))
                
                # 如果时间已经是24小时制（>=12），不需要加偏移
                # 只在12小时制时间（<12）时根据上午/下午/晚上调整
                if '下午' in row_content and start_hour < 12:
                    start_hour += 12
                    end_hour += 12
                elif '晚上' in row_content and start_hour < 12:
                    start_hour += 12
                    end_hour += 12
                # 上午时间保持不变
                
                start_time = f"{start_hour:02d}:{start_minute:02d}:00"
                end_time = f"{end_hour:02d}:{end_minute:02d}:00"
                print(f"[DEBUG] Time range found: {start_time} - {end_time}")
            else:
                # 支持全角冒号和半角冒号
                time_match = re.search(r'(\d{1,2})[:：]\s*(\d{2})', row_content)
                if time_match:
                    hour = int(time_match.group(1))
                    # 如果时间已经是24小时制（>=12），不需要加偏移
                    if '下午' in row_content and hour < 12:
                        hour += 12
                    elif '晚上' in row_content and hour < 12:
                        hour += 12
                    start_time = f"{hour:02d}:{time_match.group(2)}:00"
                    end_time = f"{hour + 1:02d}:{time_match.group(2)}:00"
                    print(f"[DEBUG] Single time found: {start_time}")
                else:
                    print(f"[DEBUG] No time found in row: {row_content}")
            
            location = ""
            title = ""
            
            for part in row_parts:
                part = part.strip()
                if not part or part.isdigit():
                    continue
                if re.search(r'[a-zA-Z]+\d+(-\d+)?', part):
                    location = part
                elif not title or len(title) < len(part):
                    title_candidate = part.replace('上午', '').replace('下午', '').replace('晚上', '').strip()
                    if title_candidate and not re.match(r'^\d+', title_candidate):
                        title = title_candidate
            
            if not title:
                title = "任务"
            
            start_time_full = f"{parsed_date}T{start_time}"
            end_time_full = f"{parsed_date}T{end_time}"
            
            items.append({
                'type': 'task',
                'user_id': user_id,
                'title': title[:50],
                'description': row_content,
                'start_time': start_time_full,
                'end_time': end_time_full,
                'location': location,
                'priority': 2,
                'category': 'work'
            })
            print(f"[DEBUG] Created task from table: title={title}, date={parsed_date}, time={start_time}-{end_time}, location={location}")
        
        if len(items) > 0:
            print(f"[DEBUG] Total items from table: {len(items)}")
            return items
    
    date_patterns = [
        r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})',
        r'(\d{4})年(\d{1,2})月(\d{1,2})日',
        r'(\d{1,2})月(\d{1,2})日',
        r'(\d{1,2})月(\d{1,2})号',
        r'(\d{1,2})/(\d{1,2})',
        r'(\d{1,2})日'
    ]
    time_pattern = r'(\d{1,2}):\s*(\d{2})'
    location_patterns = [
        r'地点[:：]\s*([^\s，,。]+)',
        r'到\s*([^\s，,。]+)',
        r'去\s*([^\s，,。]+)',
        r'目的地[:：]\s*([^\s，,。]+)',
        r'在\s*([^\s，,。]+)',
        r'会议[室厅]\s*([^\s，,。]*)',
        r'会议室\s*(\d+)?',
        r'地址[:：]\s*([^\s，,。]+)',
        r'酒店[:：]\s*([^\s，,。]+)',
        r'(\d+楼[\s\S]*?室)'
    ]
    title_patterns = [
        r'【(.+?)】',
        r'[主题|标题][:：](.+)',
        r'事项[:：](.+)',
        r'活动[:：](.+)',
        r'内容[:：](.+)'
    ]
    
    # 智能判断类型
    item_type = 'task'  # 默认为任务
    type_score = {'trip': 0, 'task': 0}
    
    for keyword in trip_keywords:
        if keyword in content:
            type_score['trip'] += 2
    
    for keyword in task_keywords:
        if keyword in content:
            type_score['task'] += 2
    
    max_type = max(type_score, key=type_score.get)
    if type_score[max_type] > 0:
        item_type = max_type
    
    print(f"[DEBUG] Type scores: {type_score}, selected type: {item_type}")
    
    parsed_date = None
    parsed_end_date = None
    parsed_time = None
    parsed_location = None
    parsed_title = None
    parsed_budget = None
    
    for date_pattern in date_patterns:
        date_match = re.search(date_pattern, content)
        if date_match:
            try:
                groups = date_match.groups()
                if len(groups) == 3:
                    y, m, d = groups
                    if len(y) == 4:
                        # 只有当确实匹配到4位数年份时才更新
                        year = int(y)
                        print(f"[DEBUG] Found explicit year: {year}")
                    else:
                        # 如果不是4位数年份，保持使用当前年份
                        month = int(y)
                        day = int(m)
                elif len(groups) == 2:
                    # 只有月日，使用当前年份
                    month = int(groups[0])
                    day = int(groups[1])
                    print(f"[DEBUG] Using current year {year} for date {month}/{day}")
                elif len(groups) == 1:
                    # 只有日期，使用当前年份和月份
                    day = int(groups[0])
                    month = today.month
                    print(f"[DEBUG] Using current year {year}, current month {month} for day {day}")
                
                if month < 1 or month > 12:
                    month = today.month
                if day < 1 or day > 31:
                    day = today.day
                
                parsed_date = f"{year:04d}-{month:02d}-{day:02d}"
                
                # 检查日期是否在过去，如果是，自动调整到下一年
                date_obj = datetime(year, month, day)
                if date_obj < today:
                    year += 1
                    parsed_date = f"{year:04d}-{month:02d}-{day:02d}"
                    print(f"[DEBUG] Date {year-1}-{month}-{day} is in the past, adjusted to {year}-{month}-{day}")
                
                print(f"[DEBUG] Found date: {parsed_date}")
                break
            except Exception as e:
                print(f"[ERROR] Date parsing error: {str(e)}")
    
    if parsed_date:
        time_match = re.search(time_pattern, content)
        if time_match:
            parsed_time = f"{time_match.group(1).zfill(2)}:{time_match.group(2)}:00"
            print(f"[DEBUG] Found time: {parsed_time}")
    
    for location_pattern in location_patterns:
        location_match = re.search(location_pattern, content)
        if location_match:
            loc = location_match.group(1).strip()
            if loc and len(loc) > 1 and len(loc) < 50:
                parsed_location = loc
                print(f"[DEBUG] Found location: {parsed_location}")
                break
    
    for title_pattern in title_patterns:
        title_match = re.search(title_pattern, content)
        if title_match:
            title = title_match.group(1).strip()
            if title and len(title) > 1:
                parsed_title = title
                print(f"[DEBUG] Found title: {parsed_title}")
                break
    
    if not parsed_title:
        all_keywords = trip_keywords + task_keywords
        for keyword in all_keywords:
            if keyword in content:
                parts = content.split(keyword)
                if len(parts) > 1:
                    parsed_title = keyword + parts[1][:10]
                else:
                    parsed_title = keyword + "安排"
                break
    
    if not parsed_title:
        clean_content = content.replace('\n', '').replace('\r', '').strip()
        parsed_title = clean_content[:30]
    
    if item_type == 'trip':
        if parsed_date:
            start_date = parsed_date
            all_dates = re.findall(r'(\d{1,2})月(\d{1,2})日|(\d{1,2})/(\d{1,2})|(\d{4})[-/](\d{1,2})[-/](\d{1,2})', content)
            if len(all_dates) > 1:
                try:
                    for date_tuple in all_dates[1:]:
                        if date_tuple[0] and date_tuple[1]:
                            end_month = int(date_tuple[0])
                            end_day = int(date_tuple[1])
                            parsed_end_date = f"{year:04d}-{end_month:02d}-{end_day:02d}"
                            break
                except:
                    parsed_end_date = (datetime.strptime(parsed_date, '%Y-%m-%d') + timedelta(days=3)).strftime('%Y-%m-%d')
            else:
                parsed_end_date = (datetime.strptime(parsed_date, '%Y-%m-%d') + timedelta(days=3)).strftime('%Y-%m-%d')
        else:
            start_date = (today + timedelta(days=7)).strftime('%Y-%m-%d')
            parsed_end_date = (today + timedelta(days=10)).strftime('%Y-%m-%d')
        
        budget_match = re.search(r'预算[:：]\s*(\d+)|费用[:：]\s*(\d+)|金额[:：]\s*(\d+)', content)
        if budget_match:
            parsed_budget = float(budget_match.group(1) or budget_match.group(2) or budget_match.group(3))
        
        items.append({
            'type': 'trip',
            'user_id': user_id,
            'title': parsed_title[:50] if parsed_title else '未命名行程',
            'description': content.replace('\n', ' ').strip(),
            'destination': parsed_location if parsed_location else '',
            'start_date': start_date,
            'end_date': parsed_end_date,
            'budget': parsed_budget
        })
        print(f"[DEBUG] Created trip: title={parsed_title}, destination={parsed_location}, dates={start_date}~{parsed_end_date}")
    
    else:
        # 任务类型（默认）
        if parsed_date:
            start_time = f"{parsed_date}T{parsed_time if parsed_time else '09:00:00'}"
            end_time = f"{parsed_date}T{parsed_time if parsed_time else '10:00:00'}"
            if parsed_time:
                try:
                    hours = int(parsed_time.split(':')[0])
                    end_time = f"{parsed_date}T{(hours + 1):02d}:00:00"
                except:
                    end_time = f"{parsed_date}T10:00:00"
        else:
            start_time = datetime.now().isoformat()
            end_time = (datetime.now() + timedelta(hours=1)).isoformat()
        
        items.append({
            'type': 'task',
            'user_id': user_id,
            'title': parsed_title[:50] if parsed_title else '未命名任务',
            'description': content.replace('\n', ' ').strip(),
            'start_time': start_time,
            'end_time': end_time,
            'location': parsed_location if parsed_location else '',
            'priority': 2,
            'category': 'work'
        })
        print(f"[DEBUG] Created task: title={parsed_title}, date={parsed_date}, location={parsed_location}")
    
    print(f"[DEBUG] Total items parsed: {len(items)}")
    return items

@app.route('/api/upload/parse-text', methods=['POST'])
def parse_text():
    try:
        data = request.get_json()
        if not data or 'text' not in data or 'user_id' not in data:
            return jsonify({'error': '缺少必要参数'}), 400
        
        user_id = data['user_id']
        text = data['text']
        
        parsed_items = parse_text_content(text, user_id)
        
        return jsonify({
            'message': '文字解析成功',
            'text': text,
            'items': parsed_items,
            'items_found': len(parsed_items) > 0
        }), 200
    except Exception as e:
        print(f"[ERROR] Text parsing failed: {str(e)}")
        return jsonify({'error': '文字解析失败: ' + str(e)}), 500

@app.route('/api/upload/image', methods=['POST'])
def upload_image():
    try:
        data = request.get_json()
        if not data or 'image' not in data or 'user_id' not in data:
            return jsonify({'error': '缺少必要参数'}), 400
        
        user_id = data['user_id']
        image_data = data['image']
        
        if image_data.startswith('data:image'):
            image_data = image_data.split(',')[1]
        
        try:
            image_bytes = base64.b64decode(image_data)
        except Exception as e:
            return jsonify({'error': '图片数据格式错误'}), 400
        
        try:
            image = Image.open(BytesIO(image_bytes))
        except Exception as e:
            return jsonify({'error': '无法解析图片文件'}), 400
        
        processed_image = preprocess_image(image)
        
        text = extract_text_from_image(processed_image)
        
        if not tesseract_found:
            return jsonify({
                'message': '图片处理完成',
                'extracted_text': '',
                'items': [],
                'text_found': False,
                'tesseract_available': False,
                'suggestion': 'OCR引擎未安装，请手动输入图片中的文字内容'
            }), 200
        
        if not text or not text.strip():
            return jsonify({
                'message': '图片处理完成',
                'extracted_text': '',
                'items': [],
                'text_found': False,
                'tesseract_available': True,
                'suggestion': '未识别到图片中的文字，请检查图片清晰度或手动输入'
            }), 200
        
        parsed_items = parse_text_content_with_ai(text, user_id)
        
        return jsonify({
            'message': '图片处理成功',
            'extracted_text': text,
            'items': parsed_items,
            'text_found': len(text.strip()) > 0,
            'tesseract_available': True
        }), 200
    except Exception as e:
        print(f"[ERROR] Image processing failed: {str(e)}")
        suggestion_text = '图片处理失败，请重试' if tesseract_found else 'OCR引擎未安装，请手动输入图片中的文字内容'
        return jsonify({
            'error': '图片处理失败: ' + str(e),
            'tesseract_available': tesseract_found,
            'suggestion': suggestion_text
        }), 500

def preprocess_image(image):
    try:
        width, height = image.size
        print(f"[DEBUG] Original image size: {width}x{height}")
        
        if width > 2000 or height > 2000:
            image = image.resize((width // 2, height // 2))
        elif width < 300 or height < 100:
            image = image.resize((width * 2, height * 2))
        
        image = image.convert('L')
        
        pixels = list(image.getdata())
        avg_brightness = sum(pixels) / len(pixels)
        print(f"[DEBUG] Average brightness: {avg_brightness}")
        
        if avg_brightness > 200:
            threshold = 200
        elif avg_brightness < 50:
            threshold = 30
        else:
            threshold = 128
        
        image = image.point(lambda p: p > threshold and 255)
        
        print("[DEBUG] Using simple PIL preprocessing")
        return image
    except Exception as e:
        print(f"[ERROR] Image preprocessing failed: {str(e)}")
        return image

def extract_text_from_image(image):
    text = ""
    
    try:
        print(f"[DEBUG] Tesseract path: {pytesseract.pytesseract.tesseract_cmd}")
    except Exception as e:
        print(f"[DEBUG] Cannot get tesseract path: {str(e)}")
    
    try:
        languages = ['chi_sim+eng', 'chi_sim']
        configs = [
            r'--oem 3 --psm 6 -c preserve_interword_spaces=1',
            r'--oem 3 --psm 3',
            r'--oem 3 --psm 11',
            r'--oem 1 --psm 6',
        ]
        
        best_text = ""
        best_length = 0
        
        for lang in languages:
            for config in configs:
                try:
                    print(f"[DEBUG] Trying language: {lang}, config: {config}")
                    current_text = pytesseract.image_to_string(image, lang=lang, config=config)
                    
                    if current_text and current_text.strip():
                        clean_text = current_text.replace('\n', '').replace('\r', '').strip()
                        print(f"[DEBUG] Found text (len={len(clean_text)}): {repr(clean_text[:50])}...")
                        
                        if len(clean_text) > best_length:
                            best_length = len(clean_text)
                            best_text = clean_text
                    else:
                        print(f"[DEBUG] No text found with lang: {lang}, config: {config}")
                except Exception as e:
                    print(f"[ERROR] Tesseract error with lang {lang}, config {config}: {str(e)}")
                    continue
        
        print(f"[DEBUG] Selected best result (len={len(best_text)}): {repr(best_text)}")
        return best_text
    except Exception as e:
        print(f"[ERROR] Tesseract not available: {str(e)}")
        return ""

@app.route('/api/ai/save_suggestion', methods=['POST'])
def save_suggestion():
    try:
        data = request.get_json()
        if not data or 'user_id' not in data or 'type' not in data:
            return jsonify({'error': 'Missing required fields'}), 400
        
        user_id = data['user_id']
        item_type = data['type']
        
        if item_type == 'task':
            new_item = Task(
                user_id=user_id,
                title=data.get('title', '未命名任务'),
                description=data.get('description'),
                start_time=datetime.fromisoformat(data.get('start_time', datetime.now().isoformat())),
                end_time=datetime.fromisoformat(data.get('end_time', (datetime.now() + timedelta(hours=1)).isoformat())),
                location=data.get('location'),
                priority=data.get('priority', 2),
                status='pending',
                category=data.get('category', 'personal')
            )
            db.session.add(new_item)
            db.session.commit()
            return jsonify({'message': 'Task saved', 'id': new_item.id}), 201
        
        elif item_type == 'trip':
            new_item = Trip(
                user_id=user_id,
                title=data.get('title', '未命名行程'),
                description=data.get('description'),
                destination=data.get('destination', ''),
                start_date=datetime.fromisoformat(data.get('start_date', datetime.now().date().isoformat())).date(),
                end_date=datetime.fromisoformat(data.get('end_date', (datetime.now() + timedelta(days=3)).date().isoformat())).date(),
                budget=data.get('budget'),
                status='planned'
            )
            db.session.add(new_item)
            db.session.commit()
            return jsonify({'message': 'Trip saved', 'id': new_item.id}), 201
        
        else:
            return jsonify({'error': 'Invalid item type'}), 400
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(host='0.0.0.0', port=5000, debug=True)